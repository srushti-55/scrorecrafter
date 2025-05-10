import re
from typing import Callable, Dict, Tuple, Any, Optional
from docx import Document
import pdfplumber
import logging

logging.getLogger("pdfminer").setLevel(logging.ERROR)

PDF_FILTERS = [
    r"#|\$|\*|\(|\)|\[|\]|%|\.|\,|/\d\d\d",
    r"COURSE NAME ISE ESE TOTAL TW PR OR TUT Tot Crd Grd GP CP PR ORD"
]
PDF_SUBSTITUTIONS = {
    r'\sAB\s': ' -- ',
    r'\sAB\s': ' -- ',
    r'\s--\s': ' --- ',
    r'\s--\s': ' --- '
}
PDF_PATTERNS = {
    'name': r"NAME\s:\s([A-Z]+\s[A-Z]+\s[A-Z]+)",
    'course': r"^(\d+[A-Z]?)\s+([A-Z: &]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z+]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)"
}

SCORE_TYPES = ['ISE', 'ESE', 'TOTAL', 'TW', 'PR', 'OR', 'TUT', 'Tot', 'Crd', 'Grd', 'GP', 'CP', 'PAR', 'ORD']

def pdf_parser(file_path: str,
               filters: list[str],
               patterns: Dict[str, str],
               substitutions: Dict[str, str],
               add_callback: Callable[[str, Any], None],
               progress_callback: Optional[Callable[[int], None]] = None) -> None:
    """
    Parse the PDF file specified by file_path.
    
    Applies filters, substitutions and regex patterns on each line in all pages.
    The extracted data is provided to add_callback.
    Optionally, progress_callback is called with the current progress percentage.
    """
    try:
        file = pdfplumber.open(file_path)
    except Exception as e:
        logging.error(e)
        return

    for page_no, page in enumerate(file.pages, start=0):
        data = page.extract_text()

        # Apply filters to remove unwanted characters or patterns
        for regex_filter in filters:
            data = re.sub(regex_filter, '', data)
        
        # Apply substitutions to standardize data
        for sub, val in substitutions.items():
            data = re.sub(sub, val, data)
        
        lines = data.splitlines()
        lines_count = len(lines)
        # Process each line and search for patterns
        for line_no, line in enumerate(lines, start=1):
            for key, pattern in patterns.items():
                match = re.search(pattern, line)
                if match:
                    # Check number of groups to determine the data structure
                    if len(match.groups()) == 1:
                        add_callback(key, match.group(1))
                    else:
                        add_callback(key, tuple(val.strip() for val in match.groups()))
            # Update progress if callback provided
            if progress_callback:
                progress_callback(int((page_no * line_no) / (len(file.pages) * lines_count) * 100))

    file.close()

def export_to_docx(file_path: str,
                   result: Dict[str, Dict[str, Dict[str, str]]],
                   options: list[Tuple[str, str]],
                   progress_callback: Optional[Callable[[float], None]] = None) -> None:
    """
    Export the result data into a docx file.
    
    Creates a table where the first column is the name and the following columns
    represent each selected course and score type. Progress is reported using progress_callback.
    """
    doc = Document()
    table = doc.add_table(rows=len(result.keys()) + 1, cols=len(options) + 1)

    # Set header for name
    table.cell(0, 0).text = 'NAME'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    # Set header for courses and score types
    for i, (course, score_type) in enumerate(options, start=1):
        table.cell(0, i).text = f'{course} ({score_type})'
        table.cell(0, i).paragraphs[0].runs[0].bold = True

    progress_tot = len(result.items()) - 1
    for i, (name, courses) in enumerate(result.items(), start=1):
        table.cell(i, 0).text = name
        # Fill the table with course scores if available
        for j, (course, score_type) in enumerate(options, start=1):
            val = courses.get(course, None)
            if val:
                table.cell(i, j).text = val[score_type]
        
        # Report progress if callback provided
        if progress_callback and progress_tot > 0:
            progress_callback((i / progress_tot) * 100)
    
    doc.save(file_path)
